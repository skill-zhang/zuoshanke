"""📝 Word 文档生成器 — 从 Markdown 或结构化数据生成 .docx

支持标题、段落、列表（有序/无序）、表格、粗体/斜体/下划线。
输出到 zuoshanke 项目下的 output/ 目录。

## 用法
    from tools.gen_docx import gen_docx
    r = json.loads(gen_docx("# 标题\\n\\n正文内容", title="报告"))
"""

import json
import os
import re
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    Document = None

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "output"
OUTPUT_DIR.mkdir(exist_ok=True)


def _add_run(paragraph, text: str, bold=False, italic=False, color=None, size=None):
    """添加 run 到段落"""
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = RGBColor(*color)
    if size:
        run.font.size = Pt(size)
    return run


def _parse_md_line(line: str, paragraph):
    """解析行内 Markdown 格式（粗体、斜体、代码、链接）"""
    # 处理加粗 **text**
    parts = re.split(r"(\*\*.+?\*\*)", line)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            _add_run(paragraph, part[2:-2], bold=True)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            _add_run(paragraph, part[1:-1], italic=True)
        else:
            # 行内代码灰底
            code_parts = re.split(r"(`[^`]+`)", part)
            for cp in code_parts:
                if cp.startswith("`") and cp.endswith("`"):
                    r = _add_run(paragraph, cp[1:-1], size=9)
                    r.font.color.rgb = RGBColor(200, 50, 50)
                else:
                    paragraph.add_run(cp)


def _md_to_docx(doc, md_text: str):
    """将 Markdown 文本写入 Document"""
    lines = md_text.split("\n")
    in_code = False
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 代码块
        if stripped.startswith("```"):
            in_code = not in_code
            i += 1
            continue

        if in_code:
            p = doc.add_paragraph()
            p.style = doc.styles["No Spacing"]
            run = p.add_run(stripped)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(50, 50, 50)
            p.paragraph_format.left_indent = Cm(1)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        # 标题
        h = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if h:
            level = len(h.group(1))
            text = h.group(2)
            heading = doc.add_heading(text, level=level)
            # 为中文标题添加中文字体
            for run in heading.runs:
                run.font.name = "微软雅黑"
                r = run._element.rPr
                if r is None:
                    r = run._element.makeelement(qn("w:rPr"), {})
                    run._element.insert(0, r)
                rFonts = r.find(qn("w:rFonts"))
                if rFonts is None:
                    rFonts = r.makeelement(qn("w:rFonts"), {})
                    r.insert(0, rFonts)
                rFonts.set(qn("w:eastAsia"), "微软雅黑")
            i += 1
            continue

        # 无序列表
        ul = re.match(r"^[\s]*[-*+]\s+(.+)$", line)
        if ul:
            p = doc.add_paragraph(style="List Bullet")
            _parse_md_line(ul.group(1), p)
            i += 1
            continue

        # 有序列表
        ol = re.match(r"^[\s]*(\d+)\.\s+(.+)$", line)
        if ol:
            p = doc.add_paragraph(style="List Number")
            _parse_md_line(ol.group(2), p)
            i += 1
            continue

        # 表格
        if stripped.startswith("|") and stripped.endswith("|"):
            rows_data = []
            while i < len(lines):
                cell_line = lines[i].strip()
                if not (cell_line.startswith("|") and cell_line.endswith("|")):
                    break
                cells = [c.strip() for c in cell_line.strip("|").split("|")]
                rows_data.append(cells)
                i += 1
            # 跳过表头分隔行
            if len(rows_data) > 1 and re.match(r"^[\s:-]+$", rows_data[1][0]):
                rows_data.pop(1)
            if rows_data:
                col_count = max(len(r) for r in rows_data)
                table = doc.add_table(rows=len(rows_data), cols=min(col_count, 8))
                table.style = "Light Grid Accent 1"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for ri, row_data in enumerate(rows_data):
                    for ci, cell_text in enumerate(row_data[:8]):
                        cell = table.cell(ri, ci)
                        cell.text = cell_text
                        # 表头加粗
                        if ri == 0:
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.bold = True
            doc.add_paragraph()  # 表后间距
            continue

        # 分隔线
        if re.match(r"^[-*_]{3,}$", stripped):
            doc.add_paragraph("─" * 40)
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        _parse_md_line(stripped, p)
        i += 1


def gen_docx(
    content: str = "",
    file_path: str = "",
    title: str = "文档",
    author: str = "坐山客",
    output_name: str = "",
) -> str:
    """从 Markdown 或结构化文本生成 Word (.docx) 文档

    Args:
        content: Markdown 文本内容（与 file_path 二选一）
        file_path: Markdown 文件路径（与 content 二选一）
        title: 文档标题
        author: 作者名称
        output_name: 输出文件名（不含 .docx，默认自动生成）

    Returns:
        JSON: {success, file_path, title, error?}
    """
    if Document is None:
        return json.dumps({"success": False, "error": "python-docx 未安装，请运行: pip install python-docx"})

    try:
        md_text = ""
        if file_path:
            if not os.path.exists(file_path):
                return json.dumps({"success": False, "error": f"文件不存在: {file_path}"})
            with open(file_path, "r", encoding="utf-8") as f:
                md_text = f.read()
        elif content:
            md_text = content
        else:
            return json.dumps({"success": False, "error": "请提供 content 或 file_path"})

        doc = Document()

        # 设置默认字体
        style = doc.styles["Normal"]
        style.font.name = "微软雅黑"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        # 封面
        for _ in range(6):
            doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(30, 60, 120)

        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"作者: {author}  |  日期: {datetime.now().strftime('%Y-%m-%d')}")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_page_break()

        # 正文
        _md_to_docx(doc, md_text)

        # 保存
        if not output_name:
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_title = re.sub(r"[^\w\u4e00-\u9fff_-]", "_", title)[:30]
            output_name = f"{safe_title}_{ts}"
        output_path = str(OUTPUT_DIR / f"{output_name}.docx")
        doc.save(output_path)

        return json.dumps({
            "success": True,
            "file_path": output_path,
            "title": title,
        }, ensure_ascii=False)

    except Exception as e:
        return json.dumps({"success": False, "error": str(e)}, ensure_ascii=False)
